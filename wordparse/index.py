from win32com import client as wc
import os
from docx import Document
from docx.shared import Inches
import re

currentDir = os.getcwd()
# 目标文件夹
targetPath = os.path.join(currentDir, 'fileTarget')

# 正则匹配
# 单、多项选择题，共计148（77 + 71）道
pattern1 = '.项选择题【.*】(\n)?'  # 选择题总共90道
pattern2 = '((单|多)项选择题【.*】\n?.*\n([A-Z]\..*\n)+)(\[答案\].*\n(\[点评\].*\n)?)?'
pattern3 = '简答题【.*】\n?(.*(\n))+\[答案\]\n?((.*（[1-9]）).*\n)+(\[点评\].*\n)?'
pattern4 = '简答题【.*】'  # 简答题总共21道
pattern5 = '综合题【.*】'  # 综合题总共11道
pattern6 = '单项选择题【.*】'  # 单选题77道
pattern7 = '多项选择题【.*】'  # 综合题总共71道

# make a new file
newQuesDoc = Document()
newAnsDoc = Document()

# step2 read file from targetPath
def readFileWithCount():
    count = 0
    def f(filepath):
        nonlocal count
        document = Document(filepath)
        totalNum = len(document.paragraphs)
        docString = ''
        for i in range(0, totalNum):
            docString += document.paragraphs[i].text + "\n"
        docString = re.sub('(\n){2,}', '', docString)
        docString.replace("脉络和复习", "")
        
        result = re.search(pattern2, docString)
        if result:
            count += 1
            haha = "题" + str(count) + ". "
            # print(haha + result.group())
            saveFile(haha + result.group(1), newQuesDoc)
            saveFile(haha + result.group(), newAnsDoc)
    return f
    # totalNum = len(document.paragraphs)
    # startNum = 0
    # while(startNum < totalNum):
    #     text = document.paragraphs[startNum].text
    #     if re.match(r'^[单|多]项选择题', text):
    #         # 题目和答案中匹配
    #         texts = ''
    #         for i in range(1, 6):
    #             texts +=  document.paragraphs[startNum + i].text
    #         if (texts.find(words) >= 0):
    #             saveFile(document, startNum)
    #         startNum += 8
    #     else :
    #         startNum += 1

readFile = readFileWithCount()

# 循环存储
def saveFile(paraStr, targetDoc):
    paraArr = paraStr.split("\n")
    for para in paraArr:
        targetDoc.add_paragraph(para)

def eachFiles():
    pathDir = os.listdir(targetPath)
    for allDir in pathDir:
        child = os.path.join(targetPath, allDir)
        readFile(child)

eachFiles()

newQuesDoc.save("题目集锦.docx")
newAnsDoc.save("答案集锦.docx")





