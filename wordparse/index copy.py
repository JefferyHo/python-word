from win32com import client as wc
import os
from docx import Document
from docx.shared import Inches
import re

originPath = "D:\\program\\wordparse\\files\\"

targetPath = "D:\\program\\wordparse\\filesTarget\\"\

# make a new file
newDoc = Document()

# 正则匹配
# 单、多项选择题
pattern1 = '[单|多]项选择题【.*】(\s|\S)(.*(\s|\S){1,})(^[单项选择题|多项选择题|简答题|综合题|脉络和复习]))'


# step2 read file from targetPath
def readFile(filepath):
    document = Document(filepath)
    docString = document.part
    print(docString)
    # totalNum = len(document.paragraphs)
    # startNum = 0
    # while(startNum < totalNum):
    #     text = document.paragraphs[startNum].text
    #     if re.match(r'^[单|多]项选择题', text):
    #         # 题目和答案中匹配
    #         texts = ''
    #         for i in range(1, 6):
    #             texts +=  document.paragraphs[startNum + i].text
    #         if (texts.find(words) >= 0):
    #             saveFile(document, startNum)
    #         startNum += 8
    #     else :
    #         startNum += 1

# 循环存储题目
def saveFile(document, startNum):
    for i in range(8):
        if i == 0:
            newDoc.add_paragraph(document.paragraphs[startNum + i].text, style='List Bullet')
        else :
            newDoc.add_paragraph(document.paragraphs[startNum + i].text, style='Normal')
    newDoc.add_paragraph()

# 循环存储答案
# def saveAnswer(document):


def eachFiles(filepath):
    pathDir = os.listdir(filepath)
    for allDir in pathDir:
        child = os.path.join('%s%s' % (filepath, allDir))
        readFile(child)

eachFiles(targetPath)
newDoc.save("集锦.docx")





